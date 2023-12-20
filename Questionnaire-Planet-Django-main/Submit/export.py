import datetime

from debugpy._vendored.pydevd.pydevd_attach_to_process.add_code_to_python_process import IS_LINUX
from django.http import JsonResponse

from Qn.models import *
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx import *
from io import BytesIO
from django.views.decorators.csrf import csrf_exempt
from Qn.form import SurveyIdForm

import djangoProject.settings
def qn_to_docx(qn_id):
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id) + ".docx"

    # code = hash_code(str(survey.username),str(qn_id))

    # docx_title = code
    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)

    paragraph_list = []

    paragraph = document.add_paragraph().add_run(survey.description, style='Song')

    introduction = "本问卷已经收集了" + str(survey.recycling_num) + "份，共计" + str(survey.question_num) + "个问题"
    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    paragraph_list.append(paragraph)

    questions = Question.objects.filter(survey_id=survey)
    i = 1
    for question in questions:

        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        document.add_paragraph().add_run(str(i) + "、" + question.title + "(" + type_str + ")", style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in ['checkbox', 'radio']:
                # option_str += alphas[option_option] + " :  "
                option_str += "选项 " + str(num) + " :  "
                option_option += 1
                num += 1

            option_str += option.content
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path


from docx2pdf import convert


def qn_to_pdf(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    if qn.type == '2':
        document, _, docx_title, docx_path = paper_to_docx(qn_id)
    elif qn.type == '3':
        document, _, docx_title, docx_path = vote_to_docx(qn_id)
    elif qn.type == '4':
        document, _, docx_title, docx_path = signup_to_docx(qn_id)
    elif qn.type == '5':
        document, _, docx_title, docx_path = epidemic_to_docx(qn_id)
    else:
        document, _, docx_title, docx_path = qn_to_docx(qn_id)
    input_file = docx_path + docx_title
    out_file = docx_path + docx_title.replace('.docx', '.pdf')
    pdf_title = docx_title.replace('.docx', '.pdf')
    try:
        import pythoncom
        pythoncom.CoInitialize()
        convert(input_file, out_file)
    except:
        doc2pdf_linux(input_file, docx_path)

    return pdf_title


@csrf_exempt
def pdf_document(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)

            pdf_title = qn_to_pdf(qn.survey_id)
            response['filename'] = pdf_title
            response['pdf_url'] = djangoProject.settings.WEB_ROOT + "/media/Document/" + pdf_title
            # TODO: 根据实时文件位置设置url
            qn.pdf_url = response['pdf_url']
            qn.save()

            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


import xlwt
from Qn.views import KEY_STR


def write_submit_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")

    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者")
    sht1.write(0, 2, "提交时间")
    question_list = Question.objects.filter(survey_id=qn)
    question_num = len(question_list)
    i = 1

    for question in question_list:
        sht1.write(0, 2 + i, str(i) + "、" + question.title)
        i += 1

    id = 1
    for submit in submit_list:
        sht1.write(id, 0, id)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id, 1, username)
        sht1.write(id, 2, submit.submit_time.strftime("%Y-%m-%d %H:%M"))
        question_num = 1
        for question in question_list:
            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR, ';')

            sht1.write(id, 2 + question_num, answer_str)

            question_num += 1

        id += 1
    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name


@csrf_exempt
def export_excel(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            username = qn.username
            # if request.session['username'] != username:
            #     response = {'status_code': 0, 'message': '没有访问权限'}
            #     return JsonResponse(response)
            try:
                submit_list = Submit.objects.filter(survey_id=qn)
                # 找不到问卷提交
            except():
                response = {'status_code': 3, 'message': '该问卷暂无提交，无法导出'}
                return JsonResponse(response)
            if len(submit_list) == 0:
                response = {'status_code': 3, 'message': '该问卷暂无提交，无法导出'}
                return JsonResponse(response)
            if qn.type == '2':
                excel_name = write_exam_to_excel(id)
            elif qn.type == '3':
                excel_name = write_vote_to_excel(id)
            elif qn.type == '4':
                excel_name = write_signup_to_excel(id)

            elif qn.type == '5':
                excel_name = write_epidemic_to_excel(id)
            else:  # TODO 其他类型
                excel_name = write_submit_to_excel(id)

            response['excel_url'] = djangoProject.settings.WEB_ROOT + "/media/Document/" + excel_name
            qn.excel_url = response['excel_url']
            response['excel_name'] = excel_name

            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)



# 分问卷类型导出word 与pdf
def paper_to_docx(qn_id): # 考试问卷导出word
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id)+"考卷" + ".docx"

    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)


    paragraph = document.add_paragraph().add_run("考试介绍： "+survey.description, style='Song')
    sum_score = 0
    questions = Question.objects.filter(survey_id=survey)
    for question in questions:
        sum_score += question.point
    introduction = "考试须知：本考卷共计" + str(survey.question_num) + "个问题，总分共计 "+str(sum_score)+"分"

    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    str_time = survey.finished_time
    if survey.finished_time  is None:
        str_time = "暂未设定"

    warning = "此外本场考试的截止时间为：" + str(str_time) + "。注意不要在考试截止时间后提交试卷！！"
    document.add_paragraph().add_run(warning, style='Song')
    questions = [] ; question_info_list = [];
    question_list = Question.objects.filter(survey_id=survey)
    for question in question_list:
        if question.type in ['school','name','class','stuId']:
            question_info_list.append(question)
        else:
            questions.append(question)
    for question in question_info_list:
        document.add_paragraph().add_run(question.title, style='Song')
    i = 1
    for question in questions:

        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        elif type == 'judge':
            type_str = "判断题"
        document.add_paragraph().add_run(str(i) + "、" + question.title + "(" + type_str +"  "+ str(question.point)+"分 )", style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in ['checkbox', 'radio','judge']:
                option_str += alphas[option_option] + " :  "
                # option_str += "选项 " + str(num) + " :  "
                option_option += 1
                num += 1

            option_str += option.content
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path
def write_exam_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")

    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者用户名")
    sht1.write(0, 2, "提交时间")
    question_list = Question.objects.filter(survey_id=qn)
    question_info_list = []; questions = []
    paper_sum_score = 0
    for question in question_list:
        paper_sum_score += question.point
        if question.type in ['school','name','class','stuId']:
            question_info_list.append(question)
        else:# TODO 都是问题吧？
            questions.append(question)
    i = 1
    for question in question_info_list:
        sht1.write(0, 2 + i,  question.title)
        i += 1

    question_num = len(questions)
    info_num = len(question_info_list)
    # i = 1

    for question in questions:
        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        elif type == 'judge':
            type_str = "判断题"
        sht1.write(0, 2 + i, str(i-info_num) + "、" + question.title+" ("+type_str+" "+str(question.point)+"分)")
        i += 1
    sht1.write(0, 2 + i, "总分 ("+str(paper_sum_score)+"分)")
    pattern_green = xlwt.Pattern()  # Create the Pattern
    pattern_green.pattern = xlwt.Pattern.SOLID_PATTERN  # May be: NO_PATTERN, SOLID_PATTERN, or 0x00 through 0x12
    pattern_green.pattern_fore_colour = 3
    # May be: 8 through 63. 0 = Black, 1 = White, 2 = Red, 3 = Green, 4 = Blue, 5 = Yellow, 6 = Magenta, 7 = Cyan, 16 = Maroon, 17 = Dark Green, 18 = Dark Blue, 19 = Dark Yellow , almost brown), 20 = Dark Magenta, 21 = Teal, 22 = Light Gray, 23 = Dark Gray, the list goes on...
    style_green = xlwt.XFStyle()  # Create the Pattern
    # style_green.pattern = pattern_green # Add Pattern to Style
    font_green = xlwt.Font()
    font_green.colour_index = 17
    style_green.font = font_green
    pattern_red = xlwt.Pattern()  # Create the Pattern
    pattern_red.pattern = xlwt.Pattern.SOLID_PATTERN  # May be: NO_PATTERN, SOLID_PATTERN, or 0x00 through 0x12
    pattern_red.pattern_fore_colour = 2
    # May be: 8 through 63. 0 = Black, 1 = White, 2 = Red, 3 = Green, 4 = Blue, 5 = Yellow, 6 = Magenta, 7 = Cyan, 16 = Maroon, 17 = Dark Green, 18 = Dark Blue, 19 = Dark Yellow , almost brown), 20 = Dark Magenta, 21 = Teal, 22 = Light Gray, 23 = Dark Gray, the list goes on...
    style_red = xlwt.XFStyle()  # Create the Pattern
    # style_red.pattern = pattern_red  # Add Pattern to Style
    font_red = xlwt.Font()
    font_red.colour_index = 2
    style_red.font = font_red
    sht1.write(1, 0, "正确答案")
    i = 2+info_num+1
    for question in questions:
        answer_str = question.right_answer
        answer_str = answer_str.replace(KEY_STR, ';')
        sht1.write(1, i, answer_str)
        i+=1

    id = 2
    for submit in submit_list:
        sht1.write(id, 0, id-1)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id, 1, username)
        sht1.write(id, 2, submit.submit_time.strftime("%Y-%m-%d %H:%M"))
        question_num = 1
        for question in question_info_list:
            answer_str = (Answer.objects.get(submit_id=submit, question_id=question)).answer
            sht1.write(id, 2 + question_num, answer_str)
            question_num += 1
        personal_score = 0
        for question in questions:

            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR,';')
            if answer.answer == question.right_answer:
                answer.score = question.point
                style = style_green
            else:
                style = style_red
                answer.score = 0
            answer.save()
            personal_score+=answer.score

            sht1.write(id, 2 + question_num, answer_str,style)

            question_num += 1
        submit.score = personal_score
        submit.save()
        sht1.write(id, 2 + question_num, submit.score)

        id += 1
    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name
# 投票问卷导出docx
def vote_to_docx(qn_id):
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id) + ".docx"

    # code = hash_code(str(survey.username),str(qn_id))

    # docx_title = code
    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)

    paragraph_list = []

    paragraph = document.add_paragraph().add_run(survey.description, style='Song')

    introduction = "本投票问卷已经收集了" + str(survey.recycling_num) + "份，共计" + str(survey.question_num) + "个问题"
    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    paragraph_list.append(paragraph)

    questions = Question.objects.filter(survey_id=survey)
    i = 1
    for question in questions:
        vote_str = ""
        if question.isVote:
            vote_str = "投票"
        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        document.add_paragraph().add_run(str(i) + "、" + question.title + "(" + vote_str+type_str + ")", style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in ['checkbox', 'radio']:
                # option_str += alphas[option_option] + " :  "
                option_str += "选项 " + str(num) + " :  "
                option_option += 1
                num += 1

            option_str += option.content
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path

# 导出投票结果导出excel
def write_vote_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")

    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者")
    sht1.write(0, 2, "提交时间")
    question_list = Question.objects.filter(survey_id=qn)
    question_sum = len(question_list)
    i = 1

    for question in question_list:
        sht1.write(0, 2 + i, str(i) + "、" + question.title)
        i += 1

    id = 1
    for submit in submit_list:
        sht1.write(id, 0, id)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id, 1, username)
        sht1.write(id, 2, submit.submit_time.strftime("%Y-%m-%d %H:%M"))
        question_num = 1
        for question in question_list:
            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR, ';')

            sht1.write(id, 2 + question_num, answer_str)

            question_num += 1

        id += 1
    question_num = 1
    submit_num = len(submit_list)
    option_id = 0

    for question in question_list:
        if question.isVote:
            result_str = ""
            option_list = Option.objects.filter(question_id=question)
            answer_list = Answer.objects.filter(question_id=question)
            option_max_num = 0
            for option in option_list:
                option_num = 0
                content = option.content
                for answer in answer_list:
                    # if answer.answer.find(content) >= 0:
                    answer_content_list = answer.answer.split(KEY_STR)
                    if content in answer_content_list:
                        option_num += 1
                if option_num > option_max_num:
                    option_max_num = option_num
                    result_str = content
                    option_id = option.option_id
            for option in option_list:
                option_num = 0
                content = option.content
                for answer in answer_list:
                    # if answer.answer.find(content) >= 0:
                    answer_content_list = answer.answer.split(KEY_STR)
                    if content in answer_content_list:
                        option_num += 1
                if option_num == option_max_num and option_id != option.option_id:
                    option_max_num = option_num
                    result_str += ";"+ content
            sht1.write(id,2+question_num,result_str)
            sht1.write(id+1, 2 + question_num, option_max_num)
        question_num += 1


    sht1.write(id, 0, "投票最高结果")
    sht1.write(id+1, 0, "投票最高人数")
    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name

# 导出疫情打卡问卷docx
def epidemic_to_docx(qn_id):
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id) + ".docx"

    # code = hash_code(str(survey.username),str(qn_id))

    # docx_title = code
    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)

    paragraph_list = []

    paragraph = document.add_paragraph().add_run(survey.description, style='Song')

    introduction = "本疫情打卡问卷已经收集了" + str(survey.recycling_num) + "份，共计" + str(survey.question_num) + "个问题需填写"
    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    paragraph_list.append(paragraph)

    questions = Question.objects.filter(survey_id=survey)
    i = 1
    for question in questions:
        vote_str = ""
        if question.isVote:
            vote_str = "投票"
        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = ""
        elif type == 'checkbox':
            type_str = ''
        elif type == 'text':
            type_str = '（填空题）'
        elif type == 'mark':
            type_str = ''
        elif type == 'location':
            type_str = '（位置信息）'
        document.add_paragraph().add_run(str(i) + "、" + question.title+type_str ,  style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in [ 'radio']:
                # option_str += alphas[option_option] + " :  "
                # option_str += "选项 " + str(num) + " :  "
                option_str = "  ⭕  "
                option_option += 1
                num += 1
            if question.type in ['checkbox']:
                # option_str += alphas[option_option] + " :  "
                # option_str += "选项 " + str(num) + " :  "
                option_str = "  □  "
                option_option += 1
                num += 1

            option_str += option.content
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path

# 导出疫情打卡问卷结果为excel
def write_epidemic_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")

    try:
        begin_date = (Submit.objects.filter(survey_id__type='5'))[0].submit_time.date()
    except:
        begin_date = datetime.datetime.today()
    print(begin_date)
    last_date = begin_date
    end_date = datetime.datetime.today()
    day_num = 1
    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者")
    sht1.write(0, 2, "提交时间")
    question_list = Question.objects.filter(survey_id=qn)
    question_sum = len(question_list)

    i = 1
    for question in question_list:
        sht1.write(0, 2 + i, str(i) + "、" + question.title)
        i += 1
    sht1.write(1,0,begin_date.strftime("%Y-%m-%d"))
    id = 1
    daka_num = 0
    style_red = xlwt.XFStyle()  # Create the Pattern
    # style_red.pattern = pattern_red  # Add Pattern to Style
    font_red = xlwt.Font()
    font_red.colour_index = 2
    style_red.font = font_red
    for submit in submit_list:
        if submit.submit_time.date()-last_date >= datetime.timedelta(days=1):
            sht1.write(id + 3 * day_num - 2, 0, "打卡人数：")
            sht1.write(id + 3 * day_num - 2, 1, daka_num)
            sht1.write(id + 3 * day_num , 0, submit.submit_time.strftime("%Y-%m-%d"))
            day_num+= 1
            daka_num = 0
            last_date = submit.submit_time.date()
        sht1.write(id+3*day_num-2, 0, id)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id+3*day_num-2, 1, username)
        sht1.write(id+3*day_num-2, 2, submit.submit_time.strftime("%Y-%m-%d %H:%M"))
        question_num = 1
        for question in question_list:
            is_red = False

            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if (question.title == "近14日内，所接触环境和人员是否一切正常？" and answer_str == "否")or (question.title == '今日本人情况是否正常？' and answer_str == "否") :
                is_red = True
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR, ';')
            if is_red:
                sht1.write(id + 3 * day_num - 2, 2 + question_num, answer_str,style_red)
            else:
                sht1.write(id+3*day_num-2, 2 + question_num, answer_str)

            question_num += 1
        daka_num += 1
        id += 1
    sht1.write(id + 3 * day_num - 2, 0, "打卡人数：")
    sht1.write(id + 3 * day_num - 2, 1, daka_num)
    question_num = 1
    submit_num = len(submit_list)
    option_id = 0

    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name

# 导出报名问卷为docx
def signup_to_docx(qn_id):
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id) + ".docx"


    # docx_title = code
    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)

    paragraph_list = []

    paragraph = document.add_paragraph().add_run(survey.description, style='Song')

    introduction = "本报名问卷已经收集了" + str(survey.recycling_num) + "份，共计" + str(survey.question_num) + "个问题，本数据仅代表该文件被导出时的数据，请及时登录问卷星球网站查看最新数据。"
    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    paragraph_list.append(paragraph)

    questions = Question.objects.filter(survey_id=survey)
    i = 1
    for question in questions:

        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "(单选题)"
        elif type == 'checkbox':
            type_str = '(多选题)'
        elif type == 'text':
            type_str = '(填空题)'
        elif type == 'mark':
            type_str = '(评分题)'
        document.add_paragraph().add_run(str(i) + "、" + question.title +  type_str , style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in ['checkbox', 'radio']:
                # option_str += alphas[option_option] + " :  "
                option_str += "选项 " + str(num) + " :  "
                option_option += 1
                num += 1

            option_str += option.content
            if option.has_num_limit:
                option_str += "  剩余"+str(option.remain_num)
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path

# 导出报名问卷结果到excel
def write_signup_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")
    submit_num = Submit.objects.filter(survey_id=qn).count()
    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者")
    sht1.write(0, 2, "提交时间")
    sht1.write(submit_num+1,0,"选择次数")

    question_list = Question.objects.filter(survey_id=qn)
    question_num = len(question_list)
    option_now_num = 0
    i = 1

    for question in question_list:

        sht1.write(0, 2 + i+option_now_num, str(i) + "、" + question.title)
        i += 1
        if not question_is_signup(question):
            continue
        options = Option.objects.filter(question_id=question)
        this_option_num = 0
        for option in options:
            this_option_num += 1

            sht1.write(0, 2 + i + option_now_num, "选项"+str(this_option_num) + "、" + option.content)
            print(option.option_id)
            sht1.write(submit_num+1, 2 + i + option_now_num, option.num_limit-option.remain_num )
            option_now_num += 1

    id = 1
    this_option_num = 0
    option_now_num = 0
    for submit in submit_list:
        sht1.write(id, 0, id)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id, 1, username)
        sht1.write(id, 2, submit.submit_time.strftime("%Y-%m-%d %H:%M"))
        option_now_num = 0
        question_num = 1
        for question in question_list:
            is_signup = question_is_signup(question)
            options = Option.objects.filter(question_id=question)
            this_option_num = 0
            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR, ';')
            pattern_green = xlwt.Pattern()  # Create the Pattern
            pattern_green.pattern = xlwt.Pattern.SOLID_PATTERN  # May be: NO_PATTERN, SOLID_PATTERN, or 0x00 through 0x12
            pattern_green.pattern_fore_colour = 0x2A
            # May be: 8 through 63. 0 = Black, 1 = White, 2 = Red, 3 = Green, 4 = Blue, 5 = Yellow, 6 = Magenta, 7 = Cyan, 16 = Maroon, 17 = Dark Green, 18 = Dark Blue, 19 = Dark Yellow , almost brown), 20 = Dark Magenta, 21 = Teal, 22 = Light Gray, 23 = Dark Gray, the list goes on...
            style_green = xlwt.XFStyle()  # Create the Pattern
            style_green.pattern = pattern_green # Add Pattern to Style
            sht1.write(id, 2 + question_num+option_now_num, answer_str)
            if is_signup:

                for option in options:
                    this_option_num += 1
                    this_answer_list = answer.answer.split(KEY_STR)
                    if option.content in this_answer_list:
                        sht1.write(id, 2+1 +option_now_num+question_num,   option.content,style_green)
                    option_now_num += 1

            question_num += 1

        id += 1


    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name
# 判断问题是否是报名题目，在导出报名excel用到，专门展示选项
def question_is_signup(question):
    is_signup = False
    options = Option.objects.filter(question_id=question)
    for option in options:
        if option.has_num_limit:
            is_signup = True
    return is_signup

def doc2pdf_linux(docPath, pdfPath):
    """
    convert a doc/docx document to pdf format (linux only, requires libreoffice)
    :param doc: path to document
    """
    cmd = 'libreoffice7.0 --headless --invisible  --convert-to pdf:writer_pdf_Export'.split() + [docPath] + [
        '--outdir'] + [pdfPath]
    print(cmd)
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=30)
    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)
